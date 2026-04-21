import io

import docx
import PyPDF2


def extract_text(file_path):
    if file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        # Extract text from docx file
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        text_parts.append(cell_text)
        return "\n".join(text_parts)
    elif file_path.endswith(".pdf"):
        reader = PyPDF2.PdfReader(file_path)
        return "\n".join(
            page.extract_text() for page in reader.pages if page.extract_text()
        )
    else:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

def chunk_text(text, chunk_size=30000, overlap=5000):
    """
    Split text into chunks with optional overlap.

    Args:
        text (str): The text to chunk.
        chunk_size (int): The size of the chunk in characters.
        overlap (int): The number of characters to overlap between chunks.

    Returns:
        list: A list of chunks.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        # If not the last chunk, try to break at a sentence or paragraph boundary
        if end < len(text):
            # try to find a good breaking point (new line or period)
            for break_char in ["\n\n", "\n", ". "]:
                last_break = text.rfind(break_char, start, end)
                if last_break != -1:
                    end = last_break + len(break_char)
                    break

        # Add the chunk to the list
        chunks.append(text[start:end].strip())
        # Mover start position with overlap
        start = end - overlap if end < len(text) else end

    return chunks
